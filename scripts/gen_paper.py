"""生成论文 Word 文档"""
import copy, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 创建新文档 ──
doc = Document()

# 设置默认样式
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ── 页面设置：A4，常规边距 ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ── 辅助函数 ──
def set_font(run, name_cn="宋体", name_en="Times New Roman", size=Pt(12), bold=False, color=None):
    run.font.size = size
    run.bold = bold
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name_en)
    rFonts.set(qn('w:hAnsi'), name_en)
    rFonts.set(qn('w:eastAsia'), name_cn)
    rFonts.set(qn('w:cs'), name_en)
    if color:
        run.font.color.rgb = color

def add_h1(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(7)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, "黑体", "Arial", Pt(14), bold=True)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(7)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, "宋体", "Times New Roman", Pt(12), bold=True)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, "宋体", "Times New Roman", Pt(12), bold=True)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_font(run, "宋体", "Times New Roman", Pt(12))
    return p

def add_img_placeholder(doc, caption):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing = 1.5
    pf.space_before = Pt(10)
    run = p.add_run(f"[ 请在此处插入：{caption} ]")
    set_font(run, "宋体", "Times New Roman", Pt(10), color=RGBColor(0x99, 0x99, 0x99))
    p2 = doc.add_paragraph()
    pf2 = p2.paragraph_format
    pf2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2.line_spacing = 1.5
    run2 = p2.add_run(caption)
    set_font(run2, "宋体", "Times New Roman", Pt(10))
    return p

# ==================== 封面 ====================
p = doc.add_paragraph()
pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.space_after = Pt(40)
run = p.add_run("[ 请在此处插入学校Logo图片 ]")
set_font(run, "宋体", "Times New Roman", Pt(11), color=RGBColor(0x99, 0x99, 0x99))

p = doc.add_paragraph()
pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.line_spacing = 1.5
run = p.add_run("四川工商学院")
set_font(run, "楷体_GB2312", "Times New Roman", Pt(36), bold=True)

p = doc.add_paragraph()
pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.line_spacing = 1.5; pf.space_after = Pt(20)
run = p.add_run("《大模型技术与应用》")
set_font(run, "宋体", "Times New Roman", Pt(26), bold=True)

doc.add_paragraph()

# 题目：居中加粗红色
p = doc.add_paragraph()
pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
pf.space_before = Pt(10); pf.space_after = Pt(20)
run = p.add_run("基于多智能体的大学辅修学习规划与跟踪系统")
set_font(run, "宋体", "Times New Roman", Pt(16), bold=True, color=RGBColor(0xEE, 0, 0))

for label, value in [
    ("学生姓名", "_________"), ("学    号", "_________"),
    ("所在学院", "人工智能与电子工程学院"), ("专业名称", "人工智能"),
    ("班    级", "2024级___班"), ("任课教师", "_________")]:
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.line_spacing = 2.0
    run = p.add_run(f"{label}：{value}")
    set_font(run, "宋体", "Times New Roman", Pt(12))

doc.add_paragraph()
p = doc.add_paragraph(); pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("四川工商学院"); set_font(run, "宋体", "Times New Roman", Pt(12))
p = doc.add_paragraph(); pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("二〇二六年六月"); set_font(run, "宋体", "Times New Roman", Pt(12))

# ==================== 摘要页 ====================
doc.add_page_break()
p = doc.add_paragraph(); pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.space_after = Pt(10)
run = p.add_run("基于多智能体的大学辅修学习规划与跟踪系统")
set_font(run, "黑体", "Arial", Pt(14), bold=True)

p = doc.add_paragraph(); pf = p.paragraph_format; pf.line_spacing = 1.5
run = p.add_run("学生：_________  指导教师：_________")
set_font(run, "宋体", "Times New Roman", Pt(12))

p = doc.add_paragraph(); pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(10)
run = p.add_run("内容摘要："); set_font(run, "黑体", "Arial", Pt(14), bold=True)

abstract = (
    "随着高等教育普及化，越来越多大学生选择辅修第二专业以增强就业竞争力。"
    "然而，辅修教育面临课程选择盲目、主辅修时间冲突频发、学习资源分散、学业进度难以全面追踪等突出问题。"
    "传统的人工咨询指导方式效率低下，难以满足大规模学生的个性化需求。"
    "本文设计并实现了一套基于多智能体的大学辅修学习规划与跟踪系统。"
    "系统采用LangGraph状态图框架构建了包含路由、课程规划、进度追踪、冲突检测、资源推荐、"
    "报告生成、数据管理、课表调整在内的八个专业化智能体，通过条件路由实现多智能体协同工作。"
    "系统集成了检索增强生成技术，基于ChromaDB向量数据库和BAAI中文嵌入模型构建知识库，"
    "覆盖辅修专业介绍、课程详情、课后习题、课堂笔记等二十九个知识文档。"
    "同时整合网络爬虫模块，从中国大学MOOC和B站等平台采集真实课程资源链接，"
    "通过后处理过滤机制确保推荐内容的真实性和可用性。"
    "前端基于Streamlit框架构建响应式Web界面，支持对话式交互、个人资料管理、数据看板等功能。"
    "系统采用DeepSeek作为主力大语言模型，并设计了LLM三层降级策略，"
    "在模型不可用时自动回退至模板规则，确保核心功能的鲁棒性。"
    "测试结果表明，系统能够有效帮助学生科学选择辅修方案，及时发现并解决学习中的问题，"
    "提升了学业管理的效率和精准度。"
)
add_body(doc, abstract)

p = doc.add_paragraph(); pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(7)
run = p.add_run(""); set_font(run, "宋体", "Times New Roman", Pt(12))

p = doc.add_paragraph(); pf = p.paragraph_format; pf.line_spacing = 1.5
run = p.add_run("关键词："); set_font(run, "黑体", "Arial", Pt(14), bold=True)
run2 = p.add_run("多智能体  LangGraph  大语言模型  辅修规划  RAG  检索增强生成  Streamlit")
set_font(run2, "宋体", "Times New Roman", Pt(12))

# ==================== 目录页 ====================
doc.add_page_break()
p = doc.add_paragraph(); pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.space_after = Pt(14)
run = p.add_run("目      录"); set_font(run, "黑体", "Arial", Pt(14), bold=True)
add_body(doc, "[ 请在此处插入自动生成的目录 ]")

# ==================== 第1章 引言 ====================
doc.add_page_break()
add_h1(doc, "1  引言")
add_h2(doc, "1.1  研究背景与意义")

add_body(doc, (
    "近年来，我国高等教育持续扩招，大学生就业竞争日趋激烈。辅修第二专业已成为大学生"
    "拓展知识结构、提升综合竞争力的重要途径。据统计，超过百分之六十的本科院校已开设辅修专业，"
    "涵盖金融学、数据科学、法学、管理学等热门方向。然而，辅修教育的快速发展也带来了一系列"
    "亟待解决的问题。"
))

add_body(doc, (
    "首先，课程选择存在盲目性。学生在选择辅修专业时往往缺乏对自身兴趣、专业匹配度和"
    "职业发展前景的系统分析，容易受同伴影响或热点趋势左右，导致选错方向后中途退出的情况频发。"
    "其次，主辅修时间冲突问题突出。主修课程和辅修课程分别由不同院系安排，缺乏统一的协调机制，"
    "学生经常面临两门课程时间重叠或考试时间撞车的困境。第三，学习资源获取效率低下。"
    "面对薄弱课程，学生不清楚哪些习题、笔记、网课资源是可靠且有针对性的，在海量信息中"
    "筛选的成本极高。第四，学业进度跟踪手段匮乏。学生难以及时掌握各门课程的成绩变化趋势"
    "和整体学业风险状况，往往在期末考试后才意识到问题的严重性。"
))

add_body(doc, (
    "随着大语言模型技术的飞速发展，以DeepSeek、GPT、Claude为代表的预训练语言模型展现出了"
    "强大的自然语言理解和生成能力，为构建智能化的学业辅导系统提供了技术基础。"
    "同时，检索增强生成技术的成熟使得将结构化知识库与LLM深度结合成为可能，"
    "有效缓解了LLM的知识幻觉问题。LangGraph等智能体编排框架的出现，"
    "则为构建多智能体协同工作的复杂应用提供了灵活的工具支撑。"
))

add_body(doc, (
    "基于以上背景，本文设计并实现了一套基于多智能体的大学辅修学习规划与跟踪系统。"
    "系统以LangGraph为智能体编排核心，整合RAG知识库检索、网络资源爬取、LLM推理等能力，"
    "通过Streamlit构建现代化的对话式Web前端，为学生提供从辅修选择到学习跟踪的一站式智能化服务。"
    "本系统的研究意义在于：将大模型技术实际应用于高等教育场景，探索多智能体架构在学业辅导"
    "领域的工程实践路径，为同类系统的设计与开发提供可参考的技术方案。"
))

add_h2(doc, "1.2  国内外研究现状")
add_h3(doc, "1.2.1  智能教育系统的研究现状")

add_body(doc, (
    "智能教育系统是人工智能技术在教育领域的重要应用方向。早期的智能导学系统主要基于"
    "知识图谱和规则引擎，如卡内基梅隆大学开发的Cognitive Tutor，通过领域知识建模和"
    "学生认知状态追踪来实现自适应教学。近年来，深度学习技术的引入使智能教育系统在"
    "学习分析、知识追踪、自动评分等方面取得了显著进展。2023年以来，大语言模型的出现"
    "进一步推动了智能教育系统的发展。Khan Academy推出的Khanmigo基于GPT-4实现了"
    "一对一的智能辅导对话；Duolingo利用GPT-4增强其语言学习体验。国内方面，"
    "学而思、作业帮等在线教育平台也积极探索LLM在解题辅导、知识点讲解中的应用。"
    "然而，现有系统多聚焦于K-12教育或单一学科的辅导，针对大学辅修这一特定场景的"
    "多智能体综合解决方案尚不多见。"
))

add_h3(doc, "1.2.2  多智能体系统的研究现状")

add_body(doc, (
    "多智能体系统通过多个专业化Agent的协作来完成复杂任务。2024年以来，随着LangGraph、"
    "AutoGen、CrewAI等框架的快速发展，基于LLM的多智能体系统研究进入活跃期。"
    "Microsoft的AutoGen框架支持多个对话Agent之间的灵活通信和任务委派；"
    "CrewAI提供了角色化的Agent协作模式。LangGraph则通过状态图模型提供了更灵活的"
    "Agent编排能力，支持条件路由、循环控制、状态持久化等高级特性。"
    "在教育领域的多智能体应用中，已有研究使用多个Agent分别负责课程推荐、学习诊断、"
    "资源检索等子任务。但多数的研究仍处于原型验证阶段，缺乏完整的工程化实现和实际部署验证。"
))

add_h3(doc, "1.2.3  RAG技术在知识密集型应用中的研究现状")

add_body(doc, (
    "检索增强生成技术通过在LLM推理前检索外部知识库，将相关文档片段注入提示词上下文，"
    "有效缓解了LLM的知识截断和幻觉问题。2023年以来，RAG技术经历了从朴素RAG到高级RAG"
    "的快速演进。朴素RAG采用简单的检索-生成串联流程，存在检索精度低、上下文散乱等问题。"
    "高级RAG通过引入查询重写、混合检索、重排序、自反思等机制显著提升了检索质量。"
    "在技术实现层面，LangChain和LlamaIndex等框架提供了完整的RAG工具链，"
    "ChromaDB、Pinecone、Weaviate等向量数据库为语义检索提供了高效的基础设施。"
    "在教育领域，RAG已被应用于智能教材问答、学习资源推荐、学业报告生成等场景。"
    "本文系统采用了基于ChromaDB的RAG架构，结合中文嵌入模型和元数据过滤机制，"
    "实现了针对辅修课程知识的高精度语义检索。"
))

add_h2(doc, "1.3  本文主要研究内容与贡献")

add_body(doc, (
    "本文围绕大学辅修学习规划与跟踪这一核心场景，研究并实现了一套基于多智能体的智能系统。"
    "主要研究内容包括：第一，设计并实现了基于LangGraph的八智能体协作架构，"
    "包含路由智能体、课程规划智能体、进度追踪智能体、冲突检测智能体、资源推荐智能体、"
    "报告生成智能体、数据管理智能体和课表调整智能体，支持七种对话阶段的复杂多轮交互。"
    "第二，构建了覆盖三个辅修专业、七门核心课程的本地知识库，包含二十九个结构化Markdown文档，"
    "通过ChromaDB向量数据库实现高效语义检索。第三，开发了网络资源爬取与验证模块，"
    "从中国大学MOOC、B站和Kaggle等平台采集真实课程链接，并通过后处理机制过滤LLM生成的虚构URL。"
    "第四，利用Streamlit框架构建了现代化的Web用户界面，集成了对话交互、个人资料管理、"
    "数据可视化看板等功能模块。"
))

add_body(doc, (
    "本文的主要贡献包括：提出了面向大学辅修场景的多智能体系统架构，实现了八种意图的"
    "智能路由与协同工作；设计了基于RAG和爬虫数据的双通道资源推荐机制，通过链接验证"
    "后处理有效抑制了LLM的知识幻觉；构建了完整的前后端工程实现并进行了端到端的功能验证，"
    "为同类系统的开发提供了可参考的实践经验。"
))

add_h2(doc, "1.4  论文组织结构")

add_body(doc, (
    "本文共分为六章。第一章为引言，介绍研究背景、国内外研究现状以及本文的主要内容和贡献。"
    "第二章介绍系统涉及的核心技术与理论基础，包括LangGraph智能体框架、RAG检索增强生成、"
    "Streamlit前端框架以及DeepSeek大语言模型等。第三章阐述系统需求分析与总体设计，"
    "包括功能性需求、系统架构设计和数据库设计。第四章详细描述系统的核心实现，"
    "包括多智能体协作流程、各Agent的具体实现以及前端界面的构建。"
    "第五章展示系统的运行效果并进行功能测试。第六章总结全文工作并展望未来的改进方向。"
))

# ==================== 第2章 ====================
doc.add_page_break()
add_h1(doc, "2  相关技术与理论基础")
add_h2(doc, "2.1  LangGraph多智能体框架")

add_body(doc, (
    "LangGraph是由LangChain团队开发的一个用于构建有状态、多角色AI应用的框架。"
    "其核心抽象是StateGraph，一种基于状态图的有向图计算模型。与传统的链式调用不同，"
    "LangGraph允许开发者在图中定义多个节点和边，通过条件路由实现复杂的控制流逻辑。"
    "每个节点可以是一个LLM调用、工具执行或任意Python函数，"
    "节点之间通过共享的State字典传递数据。"
))

add_body(doc, (
    "LangGraph的关键特性包括：持久化状态管理，每个节点的输出合并到全局State中，"
    "支持断点续传和人工介入；条件路由，基于State中的字段动态决定下一个执行的节点；"
    "循环支持，可以在图中形成循环结构以支持多轮对话；流式输出，支持节点执行过程中的"
    "流式数据推送。本系统利用LangGraph的这些特性，实现了八个Agent之间的灵活协作"
    "和七种对话阶段的状态保持，支撑了从辅修推荐到课表调整的完整业务流程。"
    "工作流图中的节点对应于各个Agent的入口函数，路由节点通过关键词加权的意图分类算法"
    "将用户请求分发到合适的Agent，冲突检测门前置检查数据完整性后再进入冲突检测Agent。"
))

add_h2(doc, "2.2  检索增强生成技术")

add_body(doc, (
    "检索增强生成是一种将信息检索与文本生成相结合的技术范式。其核心思想是在LLM生成回答之前，"
    "先从外部知识库中检索与用户查询语义相关的文档片段，将这些片段作为上下文信息注入提示词中，"
    "从而使LLM能够基于外部知识而非仅依赖模型参数中的记忆来生成回答。"
))

add_body(doc, (
    "本系统的RAG流程分为离线索引和在线检索两个阶段。离线阶段，系统扫描知识库目录下的"
    "Markdown文件，通过YAML前置元数据解析提取标题、课程代码、类型等结构化信息，"
    "将正文按四百字符的块大小进行重叠切分，然后使用BAAI的bge-small-zh-v1.5中文嵌入模型"
    "将文本块映射为384维向量，批量存入ChromaDB持久化向量数据库。在线检索阶段，"
    "用户的自然语言查询经过相同的嵌入模型向量化后，在ChromaDB中执行余弦相似度搜索，"
    "返回top-k相关文档片段，同时支持按课程代码、专业方向、资源类型等元数据字段进行过滤。"
    "检索结果经Redis缓存后与爬虫采集的外部资源链接合并，形成最终的LLM提示词上下文。"
))

add_h2(doc, "2.3  Streamlit前端框架")

add_body(doc, (
    "Streamlit是一个面向数据科学和机器学习应用的Python Web框架，"
    "其核心特点是将Python脚本直接转化为交互式Web应用，无需编写HTML、CSS或JavaScript代码。"
    "Streamlit采用声明式编程模型，通过简单的API调用即可创建文本、表格、图表、"
    "输入框、按钮等界面组件。每次用户交互时，Streamlit从顶到底重新执行完整的Python脚本，"
    "并通过内置的会话状态机制在多次执行之间保持数据。"
))

add_body(doc, (
    "本系统基于Streamlit 1.58版本构建前端界面，利用st.chat_message和st.chat_input组件"
    "实现了对话式的用户交互模式；通过st.sidebar构建了包含个人资料编辑、功能导航、"
    "系统状态指示的侧边栏面板；使用st.expander和st.tabs组织课表、进度、冲突三个Tab的"
    "数据看板；集成了Plotly Express实现学习成绩趋势图的可视化渲染。"
    "系统还利用JavaScript客户端脚本实现了LLM请求期间的加载遮罩，"
    "通过st.cache_resource实现了智能热重载机制，在代码变更时自动重建Agent实例，"
    "常规交互时直接复用缓存，显著降低了页面响应延迟。"
))

add_h2(doc, "2.4  DeepSeek大语言模型")

add_body(doc, (
    "DeepSeek是由深度求索公司开发的大语言模型系列，以其优秀的性能和极高的性价比"
    "在中文NLP领域获得了广泛关注。本系统使用DeepSeek-Chat作为主力推理模型，"
    "该模型支持128K上下文窗口，具备出色的中英文理解和生成能力，能够胜任课程推荐、"
    "资源规划、报告生成等复杂任务。同时系统设计了LLM三层降级策略：首选DeepSeek，"
    "当DeepSeek不可用时自动切换至Anthropic Claude作为备选，若两者均不可用则回退至"
    "基于规则模板的确定性输出，确保系统的核心功能在任何情况下都能正常运行。"
    "嵌入模型同样采用了双层降级设计：优先使用sentence-transformers的bge-small-zh-v1.5"
    "中文嵌入模型进行语义向量化，离线环境下自动切换至基于TF-IDF字符级哈希的降级方案，"
    "输出相同维度的向量以保持与ChromaDB的兼容性。"
))

# ==================== 第3章 ====================
doc.add_page_break()
add_h1(doc, "3  系统需求分析与总体设计")
add_h2(doc, "3.1  系统功能性需求")

add_body(doc, (
    "本系统的核心用户群体为有辅修意向或已选择辅修的大学本科生。经分析，系统的功能性需求"
    "主要包括以下八个方面。智能辅修推荐：系统根据用户的专业背景、年级、兴趣方向和课余时段，"
    "结合知识库中的辅修专业信息与报名规则，通过LLM生成个性化辅修方案并支持多轮交互确认。"
    "课表冲突检测：自动扫描主修课表与辅修课表，检测上课时间段重叠和考试时间冲突，"
    "按三个严重等级分类展示并提供解决建议。学习进度追踪：支持自然语言录入每周的作业成绩、"
    "小测成绩和考勤状态，自动计算各门课程的成绩均分、出勤率等统计指标，识别薄弱课程。"
    "个性化资源推荐：分为薄弱点补救和进阶拓展两种模式，整合RAG知识库检索和爬虫数据，"
    "推荐针对性的课后习题、课堂笔记、网课视频等学习资源。学业报告生成：汇总用户的多维学习数据，"
    "计算综合风险评分和各课程趋势，生成结构化的Markdown学业报告。"
    "数据管理：提供自然语言课表录入、退课、换课等功能，支持个人资料在线编辑。"
    "对话问答：支持学业相关的知识问答和日常闲聊。知识库管理：支持Markdown文档的"
    "自动化向量索引和增量更新。"
))

add_h2(doc, "3.2  系统架构设计")

add_body(doc, (
    "系统采用分层架构设计，自底向上分为四层。数据存储层包含SQLite关系数据库、"
    "ChromaDB向量数据库和Redis缓存。SQLite作为主存储负责持久化用户信息、课表数据、"
    "学习进度和冲突记录；ChromaDB存储知识库文档的向量嵌入，支撑语义检索；"
    "Redis提供会话状态和RAG查询结果的缓存加速，不可用时自动降级为内存字典存储。"
    "工具与基础设施层包含了嵌入模型服务、LLM调用封装、爬虫模块、统计分析工具等可复用的"
    "组件。LLM调用封装实现了多提供商支持和三层降级逻辑；爬虫模块负责从中国大学MOOC、"
    "B站等平台采集课程资源；统计分析工具提供成绩趋势、冲突汇总、风险评分等计算功能。"
    "智能体层是系统的核心，包含八个专业化Agent，每个Agent通过统一的工具接口访问底层服务。"
    "LangGraph状态图负责Agent之间的路由和协作。展示层基于Streamlit构建Web用户界面，"
    "通过Python脚本直接渲染交互式组件，与智能体层通过AgentState字典交换数据。"
))

add_img_placeholder(doc, "图3-1  系统总体架构图")

add_h2(doc, "3.3  多智能体工作流设计")

add_body(doc, (
    "系统的多智能体协作基于LangGraph的StateGraph实现。工作流包含一个路由节点、"
    "八个Agent节点和一个冲突检测门。路由节点负责对用户输入进行意图分类，"
    "采用关键词加权评分和正则匹配相结合的两阶段分类策略，识别出九种意图类型后通过"
    "条件路由分发到对应的Agent节点。冲突检测门在进入冲突检测Agent之前对数据完整性"
    "进行前置校验，确保主修课表和辅修课程已录入后才执行冲突计算，否则返回友好的提示信息。"
    "每个Agent节点执行完毕后将结果写入State字典的final_response字段，同时更新"
    "conversation_phase字段以支持多轮对话状态保持。系统支持七种对话等待阶段，"
    "分别是等待选课、等待兴趣输入、等待成绩录入、等待课表录入、等待冲突决策、"
    "等待资源反馈和等待课程调整。在等待阶段中，用户的后续输入直接路由给对应的Agent处理，"
    "无需重新进行意图分类。"
))

add_img_placeholder(doc, "图3-2  LangGraph智能体工作流图")

add_h2(doc, "3.4  数据库设计")

add_body(doc, (
    "系统使用SQLite作为主存储数据库，设计了六张核心数据表。用户表存储学生的基本信息，"
    "包括姓名、主修专业、年级、可用时段、兴趣方向、辅修状态和辅修专业名称等字段，"
    "其中可用时段和兴趣以JSON字符串格式存储以支持灵活的数据结构。课程表存储课程目录信息，"
    "包括课程代码、课程名称、学分、课程类型、所属辅修专业、前后置依赖和开设学期等字段。"
    "用户课表表记录每个学生的选课安排，包括课程代码、课程类型、上课星期、起止节次、"
    "上课地点、考试时间等字段。学习进度表追踪学生每周的学习数据，包括周次、课程代码、"
    "作业成绩、小测成绩、考勤状态等字段。冲突记录表存储检测到的时间冲突，"
    "包括冲突双方课程代码、重叠详情JSON、严重程度和建议方案等字段。"
    "报告表归档生成的学业报告，包括报告内容和生成时间等字段。"
    "数据库采用WAL模式以提升并发写入性能，启用了外键约束保证数据引用完整性。"
))

# ==================== 第4章 ====================
doc.add_page_break()
add_h1(doc, "4  系统核心实现")
add_h2(doc, "4.1  路由与意图分类模块")

add_body(doc, (
    "路由模块是系统的入口，负责将用户的自然语言输入映射到正确的Agent。"
    "系统实现了两阶段意图分类策略。第一阶段为关键词加权评分，系统维护了一个包含三十余个"
    "关键词的意图映射表，每个关键词对应一个意图类别和权重。关键词的权重取其字符长度，"
    "使得更具体的短语获得更高的匹配优先级。系统扫描用户输入中匹配到的所有关键词，"
    "按意图类别累加权重得分，最后选择得分最高的意图作为分类结果。"
    "例如，用户输入中的推荐辅修匹配plan_course意图的关键词，权重为四；"
    "冲突匹配check_conflict意图的关键词，权重为二，系统综合得分后选择最高分的意图进行路由。"
))

add_body(doc, (
    "第二阶段为LLM规则回退。当关键词匹配得分低于阈值或匹配结果模糊时，"
    "系统使用基于正则表达式的加权规则进行二次分类。同时在多轮对话场景中，"
    "如果当前conversation_phase非空闲状态，系统将跳过意图分类，直接路由到"
    "当前对话阶段对应的Agent，确保多轮交互的连续性。"
))

add_h2(doc, "4.2  资源推荐Agent的实现")

add_body(doc, (
    "资源推荐Agent是系统中逻辑最复杂的Agent，集成了薄弱点识别、RAG检索、"
    "爬虫数据整合和LLM推荐生成等多个子模块。以下详细阐述其核心实现流程。"
))

add_h3(doc, "4.2.1  薄弱点识别算法")

add_body(doc, (
    "薄弱点识别模块从SQLite读取学生的全部学习进度记录，按课程代码聚合计算每门课程的"
    "平均作业成绩、平均小测成绩和缺勤次数。识别规则如下：平均作业成绩低于七十分、"
    "平均小测成绩低于六十分、缺勤次数大于等于两次，满足任一条件的课程被标记为薄弱课程。"
    "薄弱课程按综合评分划为三个等级：缺勤两次以上或两项成绩均不达标为高风险的红色等级，"
    "一项成绩不达标为中风险的黄色等级，成绩轻微偏低为低风险的绿色等级。"
    "同时构建包含所有课程信息的课程摘要列表，供进阶拓展模式使用。"
))

add_h3(doc, "4.2.2  双通道资源检索")

add_body(doc, (
    "资源检索采用RAG知识库与网络爬虫数据并存的双通道设计。RAG通道通过Retriever对象的"
    "search_resources方法，分别检索与薄弱课程相关的习题、笔记和网课三种类型的知识库文档。"
    "检索结果按距离排序后截断至top-k，并聚合为包含exercises、notes、videos三个子列表的"
    "结构化字典。爬虫通道通过search_scraped函数对本地JSON文件进行关键词匹配搜索。"
    "系统维护了一张课程代码到中文关键词的映射表，如FIN101映射至金融学基础、金融学、"
    "货币银行、金融市场，DS201映射至机器学习、监督学习、无监督学习、深度学习，"
    "有效提升了爬虫资源的课程匹配召回率。爬取结果中的外部链接按matched_courses字段"
    "归类为课程到URL列表的映射字典。"
))

add_h3(doc, "4.2.3  LLM推荐生成与链接验证")

add_body(doc, (
    "推荐分为两种模式。薄弱点补救模式用于存在薄弱课程的情况，提示词聚焦于基础巩固，"
    "引导LLM推荐课后习题、知识点复习笔记和入门网课。进阶拓展模式用于成绩优秀的情况，"
    "提示词包含每门课程的真实可用链接列表，引导LLM生成学术前沿、行业应用、竞赛备战、"
    "项目实战和交叉学科五个方向中最为匹配的高阶推荐。两种模式的提示词均包含从RAG检索"
    "到的知识库片段和从爬虫采集的外部资源信息。"
))

add_body(doc, (
    "LLM返回JSON格式的结构化推荐列表后，系统执行链接验证后处理。系统收集全部爬虫数据库中"
    "的真实URL构建白名单集合，逐一检查LLM返回的external_link字段，将不在白名单中的URL"
    "清空，有效过滤了LLM虚构的链接。同时在格式化输出阶段，对于LLM未填写链接但存在对应"
    "爬虫数据的课程，系统从爬虫链接映射字典中强制注入一个真实链接，确保每门课程都能获得"
    "可靠的网络资源推荐。"
))

add_img_placeholder(doc, "图4-1  资源推荐Agent工作流程图")

add_h2(doc, "4.3  冲突检测Agent的实现")

add_body(doc, (
    "冲突检测Agent实现了细粒度的时间重叠检测算法。系统遍历每对主修与辅修课程的组合，"
    "首先检查两门课程是否在同一天有安排，如果同天则进一步检查上课时间段是否存在重叠。"
    "时间段重叠的计算采用区间交集算法，统计重叠的连续节次数。"
    "考试时间冲突检测支持ISO标准日期解析和模糊匹配两种方式，以两小时为冲突判定窗口。"
    "冲突按严重程度分为三个等级：考试冲突或超过两节时间段重叠为严重等级，"
    "一到两节重叠为警告等级，其他为提示等级。严重的冲突附带基于规则的解决建议，"
    "如优先调整辅修课程、咨询教务老师、申请缓修等。检测完成后将所有冲突写入"
    "conflict_history表持久化保存。"
))

add_h2(doc, "4.4  前端界面实现")

add_body(doc, (
    "前端界面基于Streamlit框架实现，整体布局为左侧边栏与右侧主聊天区的双栏结构。"
    "侧边栏顶部展示用户头像、姓名、专业和年级信息，通过Popover弹窗提供个人资料编辑表单，"
    "支持修改姓名、主修专业、年级和兴趣方向，保存后通过upsert_user写入SQLite数据库。"
    "侧边栏中部为八个功能导航按钮，点击任一按钮将对应的命令文本设为待发送消息，"
    "通过统一的pending_input机制触发Agent调用。侧边栏底部展示LLM在线状态和消息计数。"
))

add_body(doc, (
    "主聊天区顶部为渐变色系统标题，其下为JavaScript客户端加载遮罩，在用户发送消息时"
    "立即显示AI思考中的提示，避免服务器处理期间的空白等待。消息历史以聊天气泡形式渲染，"
    "用户消息和助手回复分别采用不同的样式。底部为st.chat_input输入框，支持自由文本输入"
    "和导航按钮触发的快捷命令。输入框下方的可折叠数据看板包含三个Tab页：课表Tab按周一至周日"
    "分组展示课程安排，辅修课程以黄色圆点标识；进度Tab展示成绩数据表格和Plotly折线趋势图；"
    "冲突Tab按严重等级列出冲突详情。"
))

add_img_placeholder(doc, "图4-2  系统主界面截图")
add_img_placeholder(doc, "图4-3  侧边栏个人资料编辑界面截图")
add_img_placeholder(doc, "图4-4  资源推荐结果展示截图")

# ==================== 第5章 ====================
doc.add_page_break()
add_h1(doc, "5  系统测试与效果展示")
add_h2(doc, "5.1  测试环境")

add_body(doc, (
    "系统测试环境配置如下：操作系统为Windows 11 Pro，Python版本为3.11，"
    "Streamlit版本为1.58，ChromaDB版本为0.4系列，SQLite版本为3.41。"
    "LLM服务使用DeepSeek-Chat模型，通过OpenAI兼容API接入。嵌入模型使用"
    "BAAI的bge-small-zh-v1.5，维度为384。知识库包含二十九个Markdown文件，"
    "覆盖三个辅修专业和七门核心课程。爬虫数据包含十三条经过人工验证的真实资源链接，"
    "来自中国大学MOOC、B站和Kaggle三个平台。"
))

add_h2(doc, "5.2  功能测试")

add_body(doc, (
    "系统围绕八个核心功能进行了端到端的功能测试。辅修推荐测试中，系统根据用户的主修专业、"
    "年级、课余时段和兴趣方向，通过RAG检索辅修专业信息并结合LLM生成了包含四门课程、"
    "十二学分的个性化辅修方案，用户确认后课程自动同步至课表。冲突检测测试中，"
    "系统成功识别了主修课表与辅修课表之间的时间段重叠，按严重等级分类展示并提供调整建议。"
    "进度追踪测试中，系统正确解析了自然语言录入的作业成绩和考勤数据，生成了成绩趋势折线图。"
    "资源推荐测试验证了双模式推荐的正确切换：薄弱课程场景下输出基础巩固资料，"
    "成绩优秀场景下输出进阶拓展方案，且每条推荐均包含从爬虫数据中匹配的真实外部链接。"
    "报告生成测试验证了学业报告的结构化输出，包含综合风险评分和各课程趋势分析。"
    "数据管理测试验证了课表录入、退课、换课等功能的数据一致性和外键约束正确性。"
    "对话问答测试验证了学术问答的RAG检索效果和日常闲聊的模板回复。"
    "降级测试验证了LLM不可用时系统自动切换至规则模板的功能，确保了核心功能的鲁棒性。"
))

add_img_placeholder(doc, "图5-1  辅修推荐功能测试截图")
add_img_placeholder(doc, "图5-2  资源推荐功能测试截图")
add_img_placeholder(doc, "图5-3  数据看板展示截图")

# ==================== 第6章 ====================
doc.add_page_break()
add_h1(doc, "6  总结与展望")
add_h2(doc, "6.1  工作总结")

add_body(doc, (
    "本文设计并实现了一套基于多智能体的大学辅修学习规划与跟踪系统，"
    "完成了从需求分析、系统设计到编码实现和功能测试的完整开发流程。"
    "系统以LangGraph为智能体编排框架，构建了包含路由、课程规划、进度追踪、冲突检测、"
    "资源推荐、报告生成、数据管理和课表调整在内的八个专业化Agent，"
    "覆盖了辅修教育从选课到跟踪的完整生命周期。系统整合了基于ChromaDB的RAG检索模块，"
    "对二十九个知识库文档进行了语义索引和高效检索；开发了网络爬虫模块，"
    "从多个在线教育平台采集了真实的课程资源链接，并通过后处理验证机制确保推荐内容的可靠性。"
    "前端基于Streamlit框架构建了现代化的对话式Web界面，支持多轮交互、个人资料编辑和"
    "数据可视化看板。系统设计了LLM三层降级和嵌入模型双层降级的鲁棒性策略，"
    "确保核心功能在任何硬件和网络环境下均可正常运行。"
))

add_h2(doc, "6.2  未来展望")

add_body(doc, (
    "本系统在以下方面存在进一步提升的空间。首先，在智能体协作的深度方面，"
    "当前各Agent之间的协作以串行路由为主，未来可引入多Agent并行执行和辩论机制，"
    "让多个Agent同时对同一问题给出独立分析后再综合决策，提升推荐的多样性和鲁棒性。"
    "其次，在知识库的动态更新方面，当前系统依赖离线构建的向量索引，"
    "未来可实现在线增量索引和自动过期机制，使用户在学习过程中产生的优质笔记和解题思路"
    "也能实时纳入知识库。第三，在个性化建模方面，当前系统使用规则化的薄弱点识别和"
    "关键词加权的兴趣匹配，未来可引入基于深度学习的用户画像模型，通过学习大量学生数据"
    "中的隐含模式，实现更加精准的学业风险预警和个性化推荐。第四，在系统部署方面，"
    "当前系统为单用户本地运行模式，未来可改造为多用户SaaS架构，增加用户认证、"
    "权限管理和数据隔离等功能，支持更大规模的实际应用。"
))

# ==================== 参考文献 ====================
doc.add_page_break()
add_h1(doc, "参考文献")

refs = [
    "LangChain AI. LangGraph: Build stateful, multi-actor applications with LLMs [EB/OL]. https://langchain-ai.github.io/langgraph/, 2024.",
    "DeepSeek AI. DeepSeek-Chat: A cost-effective large language model [EB/OL]. https://platform.deepseek.com/docs, 2024.",
    "ChromaDB. The open-source AI application database [EB/OL]. https://docs.trychroma.com/, 2024.",
    "Streamlit Inc. Streamlit: A faster way to build and share data apps [EB/OL]. https://docs.streamlit.io/, 2024.",
    "Lewis P, Perez E, Piktus A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks [J]. Advances in Neural Information Processing Systems, 2020, 33: 9459-9474.",
    "BAAI. BGE: BAAI General Embedding [EB/OL]. https://huggingface.co/BAAI/bge-small-zh-v1.5, 2023.",
    "Li J, Wang X, Wu Z, et al. A Survey of Large Language Models for Education [J]. arXiv preprint arXiv:2403.14832, 2024.",
    "Wu Q, Bansal G, Wang D, et al. AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation [J]. arXiv preprint arXiv:2308.08155, 2023.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing = 1.5
    run = p.add_run(f"[{i}] {ref}")
    set_font(run, "宋体", "Times New Roman", Pt(10.5))

# ── 保存 ──
out_path = "docs/基于多智能体的大学辅修学习规划与跟踪系统.docx"
doc.save(out_path)
print(f"Done: {out_path}")
print(f"Size: {os.path.getsize(out_path)} bytes")
